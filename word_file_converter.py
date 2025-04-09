from docx import Document
from docx.shared import Pt
import markdown
from bs4 import BeautifulSoup  # To parse HTML from Markdown

def md_to_docx(input_md):
    # Read Markdown file
    with open(input_md, "r", encoding="utf-8") as f:
        md_text = f.read()

    # Convert Markdown to HTML
    html_content = markdown.markdown(md_text)

    # Parse HTML to extract text and formatting
    soup = BeautifulSoup(html_content, "html.parser")

    # Create a Word document
    doc = Document()

    # Process each HTML element
    for element in soup:
        if element.name == "h1":
            doc.add_heading(element.text, level=1)
        elif element.name == "h2":
            doc.add_heading(element.text, level=2)
        elif element.name == "p":
            doc.add_paragraph(element.text)
        elif element.name == "strong" or element.name == "b":
            p = doc.add_paragraph()
            p.add_run(element.text).bold = True
        elif element.name == "em" or element.name == "i":
            p = doc.add_paragraph()
            p.add_run(element.text).italic = True
        else:
            # Fallback: Add plain text
            doc.add_paragraph(element.text)

    # Save the document
    doc.save("research.docx")
   # print(f"Converted {input_md} to {output_docx}")

# Usage
md_to_docx("research.md")